import os
import re
import csv
from docx import Document

folder = "C:/pay/spk_belgeleri"
output_csv = "C:/pay/spk_belgeleri/vakalar.csv"

skip = {"Gönderim Tarihi.docx"}

rows = []
files_read = 0
tickers_found = 0
dates_found = 0

for fname in sorted(os.listdir(folder)):
    if not fname.endswith(".docx"):
        continue
    if fname.startswith("~$"):
        continue
    if fname in skip:
        continue

    fpath = os.path.join(folder, fname)
    try:
        doc = Document(fpath)
    except Exception as e:
        print(f"HATA [{fname}]: {e}")
        continue

    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    full_text = para_text + "\n" + table_text
    # Çok satırlı eşleşmeler için tek satıra indir
    full_oneline = re.sub(r'\s+', ' ', full_text)
    files_read += 1

    # --- Ticker ---
    # 1) Köşeli parantez: [IHGZT] veya [EGEEN, JANTS, PRKAB, DOBUR]
    tickers = []
    bracket_contents = re.findall(r'\[([A-Z0-9 ,]{2,50})\]', full_text)
    for content in bracket_contents:
        for t in re.split(r'[,\s]+', content):
            t = t.strip()
            if re.match(r'^[A-Z]{2,8}[0-9]?$', t):
                tickers.append(t)

    # 2) Eğer köşeli parantezde ticker yoksa, "İlgili Şirketler" satırındaki ticker'ı ara
    if not tickers:
        ilgili = re.search(
            r'İlgili\s+Şirketler?\s*\[?([A-Z0-9, ]+)\]?', full_oneline
        )
        if ilgili:
            for t in re.split(r'[,\s]+', ilgili.group(1)):
                t = t.strip()
                if re.match(r'^[A-Z]{2,8}[0-9]?$', t):
                    tickers.append(t)

    # 3) Son çare: "(TICKER)" formatı — parantez içi 3-6 büyük harf, bilinen SPK kalıbı
    if not tickers:
        paren_tickers = re.findall(r'\(([A-Z]{3,8})\)', full_text)
        # Genel Türkçe kısaltmaları filtrele
        exclude = {"SPK", "SPKn", "AŞ", "TAŞ", "KAP", "DKB", "MKK"}
        paren_tickers = [t for t in paren_tickers if t not in exclude]
        tickers.extend(paren_tickers)

    # Tekrar edenleri kaldır, sırayı koru
    tickers = list(dict.fromkeys(tickers))

    # --- Manipülasyon dönemi ---
    period_pattern = r'(\d{2}\.\d{2}\.\d{4})\s*[-–]\s*(\d{2}\.\d{2}\.\d{4})\s+döneminde'
    periods = re.findall(period_pattern, full_oneline)

    # --- Gönderim Tarihi (veya dokümandaki ilk tarih) ---
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    karar = ""

    # 1) "Gönderim Tarihi" / "Gonderim Tarihi" etiketinin ardındaki tarih
    for i, p in enumerate(paragraphs):
        if re.search(r'[Gg][oöô0]n[de]er[iı]m\s+[Tt]arihi', p, re.IGNORECASE):
            # Etiketten sonraki ilk tarih içeren paragrafı ara (max 5 satır)
            for j in range(i + 1, min(len(paragraphs), i + 6)):
                m = re.search(r'(\d{2}\.\d{2}\.\d{4})', paragraphs[j])
                if m:
                    karar = m.group(1)
                    break
            break

    # 2) Etiket yoksa dokümandaki ilk tarih
    if not karar:
        m = re.search(r'(\d{2}\.\d{2}\.\d{4})', full_oneline)
        if m:
            karar = m.group(1)

    ticker_str = ";".join(tickers) if tickers else ""
    bas = periods[0][0] if periods else ""
    bit = periods[0][1] if periods else ""

    if ticker_str:
        tickers_found += len(tickers)
    if bas and bit:
        dates_found += 1

    print(f"[{fname}] ticker={tickers} dönem={periods[:1]} karar=['{karar}']")

    rows.append({
        "dosya": fname,
        "ticker": ticker_str,
        "baslangic": bas,
        "bitis": bit,
        "karar_tarihi": karar
    })

# CSV yaz
with open(output_csv, "w", newline="", encoding="utf-8-sig") as f:
    writer = csv.DictWriter(f, fieldnames=["dosya", "ticker", "baslangic", "bitis", "karar_tarihi"])
    writer.writeheader()
    writer.writerows(rows)

print("\n" + "="*60)
print(f"Okunan dosya sayısı  : {files_read}")
print(f"Bulunan ticker sayısı: {tickers_found}")
print(f"Tarih tespiti (dönem): {dates_found}")
print(f"CSV kaydedildi       : {output_csv}")
